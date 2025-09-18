"""
DOCX document processor using python-docx.
"""

from pathlib import Path
from typing import List, Dict, Any, Optional
import re

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

from ...models.document_chunk import DocumentChunk
from ...config.settings import get_settings


class DocxProcessor:
    """Process DOCX documents with structure preservation."""

    def __init__(self):
        self.settings = get_settings()

    async def extract_text(self, file_path: Path) -> str:
        """
        Extract text content from DOCX file with structure preservation.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content with formatting markers

        Raises:
            Exception: If DOCX cannot be processed
        """
        try:
            document = DocxDocument(str(file_path))

            content_parts = []
            current_section = None

            for element in document.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    paragraph = next((p for p in document.paragraphs if p.element == element), None)
                    if paragraph:
                        text = paragraph.text.strip()
                        if text:
                            # Detect headings based on style or formatting
                            if self._is_heading(paragraph):
                                heading_level = self._get_heading_level(paragraph)
                                if heading_level <= 3:  # Only track major headings
                                    current_section = text
                                    content_parts.append(f"\n## {text} ##\n")
                                else:
                                    content_parts.append(f"\n{text}\n")
                            else:
                                content_parts.append(text)

                elif element.tag.endswith('tbl'):  # Table
                    table_text = self._extract_table_text(element, document)
                    if table_text:
                        content_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")

            return "\n".join(content_parts)

        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX {file_path}: {str(e)}")

    async def create_chunks(self, content: str, document_id: str) -> List[DocumentChunk]:
        """
        Create text chunks from DOCX content with section tracking.

        Args:
            content: Extracted text content
            document_id: Document identifier

        Returns:
            List of document chunks with section information
        """
        chunk_size = self.settings.document.chunk_size
        chunk_overlap = self.settings.document.chunk_overlap

        chunks = []
        start_pos = 0
        chunk_index = 0
        current_section = None

        while start_pos < len(content):
            end_pos = min(start_pos + chunk_size, len(content))

            # Extract chunk content
            chunk_content = content[start_pos:end_pos]

            # Find current section for this chunk
            section_title = self._find_current_section(content, start_pos, end_pos)
            if section_title:
                current_section = section_title

            # Try to break at natural boundaries
            if end_pos < len(content):
                # Look for section breaks, paragraph breaks, then sentence endings
                search_start = max(end_pos - 300, start_pos)

                boundaries = [
                    ('##', 0),      # Section headings (highest priority)
                    ('\n\n', 2),    # Paragraph breaks
                    ('. ', 1),      # Sentence endings
                    ('! ', 1),
                    ('? ', 1),
                ]

                best_boundary = None
                best_priority = float('inf')

                for boundary, priority in boundaries:
                    pos = chunk_content.rfind(boundary, search_start - start_pos)
                    if pos > 0 and priority < best_priority:
                        best_boundary = pos + len(boundary)
                        best_priority = priority

                if best_boundary:
                    end_pos = start_pos + best_boundary
                    chunk_content = content[start_pos:end_pos]

            chunk_content = chunk_content.strip()

            if chunk_content:
                # Clean up formatting markers for display
                clean_content = self._clean_formatting_markers(chunk_content)

                if clean_content.strip():
                    chunk = DocumentChunk(
                        document_id=document_id,
                        content=clean_content,
                        start_char=start_pos,
                        end_char=end_pos,
                        chunk_index=chunk_index,
                        section_title=current_section,
                        embedding_vector=[0.0] * 384,  # Placeholder
                        metadata={
                            "processor": "docx",
                            "section_title": current_section,
                            "has_tables": "[TABLE]" in chunk_content,
                            "has_headings": "##" in chunk_content
                        }
                    )
                    chunks.append(chunk)
                    chunk_index += 1

            # Move to next chunk with overlap
            start_pos = max(end_pos - chunk_overlap, start_pos + 1)

        return chunks

    def _is_heading(self, paragraph) -> bool:
        """
        Check if paragraph is a heading based on style or formatting.

        Args:
            paragraph: python-docx paragraph object

        Returns:
            True if paragraph appears to be a heading
        """
        # Check style name
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name or 'title' in style_name:
                return True

        # Check formatting - bold, larger font, etc.
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.bold or (first_run.font.size and first_run.font.size > Inches(0.15)):
                return True

        # Check for common heading patterns
        text = paragraph.text.strip()
        if text:
            # All caps short text
            if text.isupper() and len(text) < 100:
                return True

            # Numbered headings (1. Title, 1.1 Subtitle, etc.)
            if re.match(r'^\d+\..*', text):
                return True

        return False

    def _get_heading_level(self, paragraph) -> int:
        """
        Determine heading level (1-6).

        Args:
            paragraph: python-docx paragraph object

        Returns:
            Heading level (1 = highest)
        """
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()

            # Extract level from style name
            if 'heading' in style_name:
                level_match = re.search(r'(\d+)', style_name)
                if level_match:
                    return min(int(level_match.group(1)), 6)
                return 1  # Default heading level

            if 'title' in style_name:
                return 1

        # Estimate level based on formatting
        text = paragraph.text.strip()
        if text.isupper():
            return 1

        if re.match(r'^\d+\.', text):
            # Count dots to determine level
            dots = text.count('.')
            return min(dots + 1, 6)

        return 2  # Default level

    def _find_current_section(self, content: str, start_pos: int, end_pos: int) -> Optional[str]:
        """
        Find the current section title for a text range.

        Args:
            content: Full document content
            start_pos: Start position of chunk
            end_pos: End position of chunk

        Returns:
            Section title or None
        """
        # Look backwards from start_pos to find the most recent section heading
        search_text = content[:start_pos]

        # Find all section markers before this position
        section_matches = re.findall(r'## (.+?) ##', search_text)

        if section_matches:
            return section_matches[-1].strip()  # Return the last (most recent) section

        return None

    def _clean_formatting_markers(self, text: str) -> str:
        """
        Clean formatting markers from text for display.

        Args:
            text: Text with formatting markers

        Returns:
            Cleaned text
        """
        # Remove section markers but preserve the heading text
        text = re.sub(r'## (.+?) ##', r'\1', text)

        # Clean up table markers
        text = re.sub(r'\[TABLE\]\n', '', text)
        text = re.sub(r'\n\[/TABLE\]', '', text)

        # Clean up excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)

        return text.strip()

    def _extract_table_text(self, table_element, document) -> str:
        """
        Extract text from table with basic formatting.

        Args:
            table_element: XML table element
            document: Document object

        Returns:
            Formatted table text
        """
        try:
            # Find the table object
            table = None
            for doc_table in document.tables:
                if doc_table.element == table_element:
                    table = doc_table
                    break

            if not table:
                return ""

            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_text.append(cell_text)

                if any(row_text):  # Only add non-empty rows
                    table_text.append(" | ".join(row_text))

            return "\n".join(table_text)

        except Exception:
            return ""

    def get_document_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from DOCX document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with document metadata
        """
        try:
            document = DocxDocument(str(file_path))

            # Core properties
            core_props = document.core_properties

            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "category": core_props.category or "",
                "comments": core_props.comments or "",
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "last_modified_by": core_props.last_modified_by or "",
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
            }

            return metadata

        except Exception as e:
            return {"error": f"Failed to extract metadata: {str(e)}"}

    def has_complex_formatting(self, file_path: Path) -> bool:
        """
        Check if document has complex formatting that might affect text extraction.

        Args:
            file_path: Path to DOCX file

        Returns:
            True if document has complex formatting
        """
        try:
            document = DocxDocument(str(file_path))

            # Check for complex elements
            has_tables = len(document.tables) > 0
            has_images = any(rel.reltype.endswith('image')
                           for rel in document.part.rels.values())

            # Check for text boxes, shapes, etc. (more complex detection)
            complex_elements = 0
            for paragraph in document.paragraphs:
                if paragraph.runs:
                    for run in paragraph.runs:
                        if run.element.xpath('.//w:drawing'):
                            complex_elements += 1

            return has_tables or has_images or complex_elements > 5

        except Exception:
            return False